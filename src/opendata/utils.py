import logging
import re
from pathlib import Path
import socket
import sys
import os
from typing import List, Set, Generator, Callable, Optional, Any, Tuple
from opendata.models import ProjectFingerprint


def setup_logging(level: int = logging.INFO):
    """Configures the global logging for the application."""
    logging.basicConfig(
        level=level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%H:%M:%S",
        handlers=[logging.StreamHandler(sys.stdout)],
    )


def get_resource_path(relative_path: str) -> Path:
    """
    Get absolute path to resource, works for dev, tests and for PyInstaller.
    Searches in bundled _MEIPASS, then in package dir, then relative to current file.
    """
    # 1. Check if running in a PyInstaller bundle
    if hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS) / relative_path

    # 2. Try to find relative to the package root (src/opendata or site-packages/opendata)
    try:
        import opendata

        pkg_path = Path(opendata.__file__).parent.absolute()

        # We search for the relative path inside the package
        # relative_path might be "opendata/prompts" or just "prompts"
        clean_rel = relative_path
        if relative_path.startswith("opendata/"):
            clean_rel = relative_path[9:]

        loc = pkg_path / clean_rel
        if loc.exists():
            return loc
    except Exception:
        pass

    # 3. Fallback to project root search (for development mode)
    # Get the directory where 'src' is
    base_path = Path(__file__).parent.parent.parent.absolute()

    locs = [
        base_path / relative_path,
        base_path / "src" / relative_path,
        base_path / "opendata" / relative_path,
    ]
    for loc in locs:
        if loc.exists():
            return loc

    return base_path / relative_path


def get_app_version() -> str:
    """Reads the application version from the VERSION file or package metadata."""
    # 1. Try to find VERSION file using get_resource_path (works in dev and bundle)
    try:
        version_file = get_resource_path("VERSION")
        if version_file.exists():
            return version_file.read_text(encoding="utf-8").strip()
    except Exception:
        pass

    # 2. Try package metadata (installed mode)
    try:
        from importlib.metadata import version

        return version("opendata-tool")
    except Exception:
        pass

    return "0.0.0"


def get_local_ip() -> str:
    """Returns the local IP address."""
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "127.0.0.1"


def shorten_path(path: str, max_len: int = 60) -> str:
    """Shortens a path by keeping the first two and last two components."""
    if len(path) <= max_len:
        return path
    parts = path.split("/")
    if len(parts) <= 4:
        return path
    return f"{parts[0]}/{parts[1]}/.../{parts[-2]}/{parts[-1]}"


def format_size(size_bytes: int) -> str:
    """Formats bytes into human readable string."""
    import math

    if size_bytes == 0:
        return "0 B"
    units = ("B", "KB", "MB", "GB", "TB")
    i = int(math.floor(math.log(size_bytes, 1024)))
    return f"{size_bytes / (1024**i):.1f} {units[i]}"


def format_file_list(files: List[Path], root: Path) -> str:
    """Formats a list of files with sizes for AI context."""
    lines = []
    for p in sorted(files):
        try:
            rel_p = p.relative_to(root)
            size = format_size(p.stat().st_size)
            lines.append(f"- {rel_p} ({size})")
        except Exception:
            continue
    return "\n".join(lines)


def is_path_excluded(rel_path_str: str, name: str, exclude_patterns: List[str]) -> bool:
    """Checks if a relative path string or filename is excluded by any pattern using robust pathlib matching."""
    if not exclude_patterns:
        return False

    # Use Path for robust glob matching (handles **/ and name matching correctly)
    p = Path(rel_path_str)

    for pattern in exclude_patterns:
        clean_pat = pattern.rstrip("/")
        try:
            # Path.match handles exact names (e.g. "build") and paths ("src/build")
            # as well as glob patterns ("**/*.py")
            if p.match(clean_pat):
                return True

            # Special case: **/foo should match "foo" at root
            if clean_pat.startswith("**/") and p.match(clean_pat[3:]):
                return True
        except ValueError:
            # Invalid pattern ignored
            pass
    return False


def walk_project_files(
    root: Path,
    stop_event: Optional[Any] = None,
    exclude_patterns: Optional[List[str]] = None,
) -> Generator[Tuple[Path, Any], None, None]:
    """
    Yields (Path, stat) tuples for all relevant files, skipping excluded ones.
    Optimized using os.scandir for high-performance directory traversal.
    """
    import os

    skip_dirs = {".git", ".venv", "node_modules", "__pycache__", ".opendata_tool"}
    root_str = str(root.expanduser().resolve())
    excludes = exclude_patterns or []

    import logging

    logger = logging.getLogger("opendata.utils")
    if excludes:
        logger.debug(f"Walking {root_str} with exclusions: {excludes}")

    def _walk(current_dir):
        if stop_event and stop_event.is_set():
            return

        # Fast check for .ignore file to skip entire directory tree
        if (Path(current_dir) / ".ignore").exists():
            return

        rel_dir = os.path.relpath(current_dir, root_str).replace("\\", "/")
        if rel_dir == ".":
            rel_dir = ""

        # Check if directory itself is excluded (defensive)
        if rel_dir and is_path_excluded(
            rel_dir, os.path.basename(current_dir), excludes
        ):
            return

        # Yield directory itself (with None for stat)
        yield Path(current_dir), None

        try:
            # Stream entries to handle massive directories gracefully
            with os.scandir(current_dir) as it:
                subdirs = []
                for entry in it:
                    if stop_event and stop_event.is_set():
                        return

                    if entry.name.startswith(".") or entry.is_symlink():
                        continue

                    rel_entry_path = (
                        os.path.join(rel_dir, entry.name).replace("\\", "/")
                        if rel_dir
                        else entry.name
                    )

                    # In-scan exclusion: Skip excluded files/dirs immediately
                    if is_path_excluded(rel_entry_path, entry.name, excludes):
                        continue

                    if entry.is_dir():
                        if entry.name in skip_dirs:
                            continue
                        subdirs.append(entry.path)
                    elif entry.is_file():
                        try:
                            yield Path(entry.path), entry.stat()
                        except OSError:
                            continue

                # Recurse after streaming current dir
                for subdir in subdirs:
                    yield from _walk(subdir)
        except OSError:
            return

    yield from _walk(root_str)


def scan_project_lazy(
    root: Path,
    progress_callback: Optional[Callable[[str, str, str], None]] = None,
    stop_event: Optional[Any] = None,
    exclude_patterns: Optional[List[str]] = None,
) -> Tuple[ProjectFingerprint, List[dict]]:
    """
    Scans a directory recursively. Optimized for huge datasets.
    Returns both the Fingerprint and the full file list for database indexing.
    """
    import time

    file_count = 0
    total_size = 0
    extensions = set()
    structure_sample = []
    full_inventory = []

    last_ui_update = 0
    UI_UPDATE_INTERVAL = 0.1

    for p, stat in walk_project_files(root, stop_event, exclude_patterns):
        if stat is None:  # It's a directory
            pass
        else:  # It's a file
            file_count += 1
            size = stat.st_size
            total_size += size
            suffix = p.suffix.lower()
            extensions.add(suffix)

            rel_path = str(p.relative_to(root))
            full_inventory.append(
                {"path": rel_path, "size": size, "mtime": stat.st_mtime}
            )

            if len(structure_sample) < 50:
                structure_sample.append(rel_path)

        now = time.time()
        if progress_callback and (now - last_ui_update > UI_UPDATE_INTERVAL):
            total_size_str = format_size(total_size)
            progress_callback(
                f"{total_size_str} - {file_count} files",
                str(p.relative_to(root)),
                f"Scanning {p.name}...",
            )
            last_ui_update = now

    fp = ProjectFingerprint(
        root_path=str(root),
        file_count=file_count,
        total_size_bytes=total_size,
        extensions=list(extensions),
        structure_sample=structure_sample,
    )
    return fp, full_inventory


def list_project_files_full(
    root: Path,
    stop_event: Optional[Any] = None,
    exclude_patterns: Optional[List[str]] = None,
) -> List[dict]:
    """
    Returns a full list of files with metadata (size, rel_path) for UI inventory.
    Skips directories, yields only files.
    """
    files = []
    for p, stat in walk_project_files(
        root, stop_event, exclude_patterns=exclude_patterns
    ):
        # walk_project_files yields (Path, stat)
        # if stat is None, it's a directory
        if stat is not None:
            rel_path = str(p.relative_to(root))
            try:
                files.append(
                    {
                        "path": rel_path,
                        "size": stat.st_size,
                        "mtime": stat.st_mtime,
                    }
                )
            except Exception:
                pass
    return files


def read_file_header(p: Path, max_bytes: int = 4096) -> str:
    """
    Reads only the first few KB of a file to detect metadata/headers.
    Safe for TB-scale data files."""
    try:
        with open(p, "rb") as f:
            chunk = f.read(max_bytes)
            # Try decoding as UTF-8, fallback to simple representation
            return chunk.decode("utf-8", errors="replace")
    except Exception:
        return ""


class PromptManager:
    """Manages external Markdown-based prompt templates."""

    def __init__(self, prompts_dir: Path | None = None):
        if not prompts_dir:
            # Use the robust resource path helper
            prompts_dir = get_resource_path("opendata/prompts")
        self.prompts_dir = prompts_dir

    def render(self, template_name: str, context: dict) -> str:
        """Loads a .md template and renders it with the provided context."""
        template_path = self.prompts_dir / f"{template_name}.md"
        if not template_path.exists():
            return f"Error: Template {template_name} not found. (Searched in {self.prompts_dir})"

        with open(template_path, "r", encoding="utf-8") as f:
            template = f.read()

        return template.format(**context)


class FullTextReader:
    """
    Utility for deep reading of research documents (LaTeX, Docx) to support
    full-text context extraction.
    """

    @staticmethod
    def read_latex_full(filepath: Path) -> str:
        r"""
        Reads the full content of a LaTeX file.
        Recursively resolves \input{} and \include{} commands.
        """
        try:
            content = []
            base_dir = filepath.parent

            def resolve_recursive(current_path: Path):
                with open(current_path, "r", encoding="utf-8", errors="replace") as f:
                    for line in f:
                        # Simple regex for \input{file} and \include{file}
                        match = re.search(r"\\(?:input|include)\{([^}]+)\}", line)
                        if match:
                            sub_file = match.group(1)
                            if not sub_file.endswith(".tex"):
                                sub_file += ".tex"
                            sub_path = base_dir / sub_file
                            if sub_path.exists():
                                resolve_recursive(sub_path)
                            else:
                                content.append(f" [Missing file: {sub_file}] ")
                        else:
                            content.append(line)

            resolve_recursive(filepath)
            return "".join(content)
        except Exception as e:
            return f"[Error reading LaTeX file: {e}]"

    @staticmethod
    def read_docx_full(filepath: Path) -> str:
        """
        Reads the full text content of a .docx file, including paragraphs and tables.
        """
        try:
            from docx import Document  # type: ignore

            doc = Document(str(filepath))
            full_text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Extract tables (naive linear reading)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n\n".join(full_text)
        except Exception as e:
            return f"[Error reading Docx file: {e}]"

    @staticmethod
    def read_ipynb_full(filepath: Path) -> str:
        """
        Reads a Jupyter Notebook (.ipynb) and returns only text from markdown and code cells.
        Skips binary outputs and metadata.
        """
        try:
            import json

            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                nb = json.load(f)

            content = []
            for cell in nb.get("cells", []):
                cell_type = cell.get("cell_type")
                source = cell.get("source", [])
                if isinstance(source, list):
                    source = "".join(source)

                if cell_type == "markdown":
                    content.append(source)
                elif cell_type == "code":
                    content.append(f"```python\n{source}\n```")

            return "\n\n".join(content)
        except Exception as e:
            return f"[Error reading Jupyter Notebook: {e}]"

    @staticmethod
    def read_full_text(filepath: Path) -> str:
        """
        Dispatches to the appropriate reader based on file extension.
        """
        suffix = filepath.suffix.lower()
        if suffix == ".tex":
            return FullTextReader.read_latex_full(filepath)
        elif suffix == ".docx":
            return FullTextReader.read_docx_full(filepath)
        elif suffix == ".ipynb":
            return FullTextReader.read_ipynb_full(filepath)
        else:
            # Fallback for plain text files
            try:
                with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                    return f.read()
            except Exception as e:
                return f"[Error reading text file: {e}]"
